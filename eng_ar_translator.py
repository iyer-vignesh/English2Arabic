import docx
import PyPDF2
from translate import Translator
from docx2pdf import convert
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st

import pythoncom

pythoncom.CoInitialize()

st.write("""
# Language Translation App

This app translates an uploaded file of **English** to **Arabic** language.
\n

Supported file formats are **.docx , .txt and .pdf files only !**
""")

uploaded_file = st.file_uploader("Choose a file", type=['docx', 'pdf', 'txt'])


def readtxt(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


def readpdf(pdfName):
    read_pdf = PyPDF2.PdfReader(pdfName)
    fulltext = ''

    for i in range(0, len(read_pdf.pages)):
        page = read_pdf.pages[i]
        page_content = page.extract_text()
        fulltext += page_content
    return fulltext


translator = Translator(from_lang='en', to_lang="ar")


def trans(text):
    marker_1 = True
    marker_2 = True

    doc = text.split('\n')  # Split text paragraph wise & have them as list
    para = [i.split('.') for i in doc]  # Sentence tokenize each paragraph & keep them in separate lists

    corpus = []

    for sentences in para:

        multi_sent = []
        single_sent = []

        if len(sentences) > 1:

            for i in sentences:

                if len(i) > 499:
                    ind = len(i) // 2
                    while marker_1 == True:
                        ind += 1
                        if i[ind] == ' ':
                            part_1 = i[:ind + 1]
                            part_2 = i[ind:]
                            marker_1 = False

                    sent_1 = translator.translate(part_1)
                    sent_2 = translator.translate(part_2)
                    sent_3 = sent_1 + ' ' + sent_2
                    multi_sent.append(sent_3)


                else:
                    txt = translator.translate(i)
                    multi_sent.append(txt)

            corpus.append(multi_sent)

        elif len(sentences) == 1:

            sent_0 = sentences[0]

            if len(sent_0) > 499:
                ind = len(sent_0) // 2
                while marker_2 == True:
                    ind += 1
                    if i[ind] == ' ':
                        part_0_1 = sent_0[:ind + 1]
                        part_0_2 = sent_0[ind:]
                        marker_2 = False

                sent_0_1 = translator.translate(part_0_1)
                sent_0_2 = translator.translate(part_0_2)
                sent_0_3 = sent_0_1 + ' ' + sent_0_2
                single_sent.append(sent_0_3)


            else:
                txt_0 = translator.translate(sent_0)
                single_sent.append(txt_0)

            corpus.append(single_sent)


        else:
            continue

    return corpus


def write_file(text):
    trans_text = []
    translated_file = 'Translated_file.docx'
    for para in text:
        temp = ''

        if len(para) > 1:
            for sent in para:
                temp += sent
            trans_text.append([temp])

        else:
            trans_text.append(para)

    # Create a new docx file
    document = docx.Document()
    # Add a heading
    # document.add_heading(trans_text[0][0], level=0)

    heading = document.add_heading(trans_text[0][0], level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for para in trans_text[1:]:
        # paragraph = document.add_paragraph(para)

        paragraph.add_run(para)

    document.save(translated_file)
    # Converting docx present in the same folder
    # as the python file
    trans_file = convert(translated_file)

    return trans_file


if uploaded_file is not None:
    st.success("File Saved")
    result = st.button(label='Submit')

    if result:

        if uploaded_file is not None:

            st.success('file Uploaded')
            file_details = {"filename": uploaded_file.name, "filetype": uploaded_file.type,
                            "filesize": uploaded_file.size}
#             st.write(file_details)

            if uploaded_file.type == "text/plain":
                # Read as string (decode bytes to string)
                raw_text = str(uploaded_file.read(), "utf-8")
#                 st.text(raw_text)
                trans_text = trans(raw_text)
                write_file(trans_text)
                st.write('Complete')


            elif uploaded_file.type == "application/pdf":
                try:
                    raw_text = readpdf(uploaded_file)
#                     st.write(type(raw_text))

                    trans_text = trans(raw_text)
#                     st.write(trans_text)
                    write_file(trans_text)
                    st.write('Complete')


                except:
                    st.write("None")

            else:

                raw_text = readtxt(uploaded_file)
#                 st.write(raw_text)
#                 st.write(type(raw_text))

                trans_text = trans(raw_text)
#                 st.write(trans_text)
                write_file(trans_text)
                st.write('Complete')

            with open("Translated_file.pdf", "rb") as pdf_file:
                PDFbyte = pdf_file.read()

            uploaded_file_name = uploaded_file.name.split('.')[0]
            new_file_name = 'Translated ' + uploaded_file_name + '.pdf'

            st.download_button(label="Download File",
                               data=PDFbyte,
                               file_name=new_file_name,
                               mime='application/octet-stream')
